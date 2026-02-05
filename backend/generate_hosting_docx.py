"""
Generate HOSTING.docx from HOSTING.md for download.
Run from project root: python backend/generate_hosting_docx.py
Output: HOSTING.docx in project root.
"""
import os
import re
import sys

# Ensure backend is on path
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def parse_md_table(lines):
    """Parse markdown table into list of rows (list of cells)."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith('#'):
            break
        if re.match(r'^[\|\s\-:]+$', line):  # separator row
            continue
        cells = [c.strip() for c in line.split('|') if c.strip() or line.count('|') > 1]
        if cells:
            rows.append(cells)
    return rows

def md_to_docx(md_path, docx_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Calibri'

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    lines = content.split('\n')
    i = 0
    in_code = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        raw = line

        if in_code:
            if line.strip().startswith('```'):
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_before = Pt(6)
                in_code = False
                code_lines = []
            else:
                code_lines.append(line)
            i += 1
            continue

        if line.strip().startswith('```'):
            in_code = True
            i += 1
            continue

        if in_table:
            if line.strip().startswith('|') and not re.match(r'^[\|\s\-:]+$', line.strip()):
                table_lines.append(line)
                i += 1
                continue
            else:
                # flush table
                if table_lines:
                    rows = parse_md_table(table_lines)
                    if rows:
                        t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        t.style = 'Table Grid'
                        for ri, row in enumerate(rows):
                            for ci, cell in enumerate(row):
                                if ci < len(t.rows[ri].cells):
                                    t.rows[ri].cells[ci].text = cell
                        doc.add_paragraph()
                in_table = False
                table_lines = []
                continue

        if line.strip().startswith('|') and not re.match(r'^[\|\s\-:]+$', line.strip()):
            in_table = True
            table_lines.append(line)
            i += 1
            continue

        if line.strip() == '---':
            i += 1
            continue

        if line.startswith('# '):
            p = doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=3)
            i += 1
            continue

        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(text)
            i += 1
            continue

        if line.strip().startswith('- [ ]'):
            text = line.strip()[5:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.add_run('☐ ' + text)
            i += 1
            continue

        if line.strip():
            p = doc.add_paragraph()
            # Inline code and bold
            rest = line.strip()
            while rest:
                m = re.search(r'\*\*(.+?)\*\*', rest)
                code = re.search(r'`([^`]+)`', rest)
                if code and (not m or code.start() < m.start()):
                    start, end = code.span()
                    if start > 0:
                        p.add_run(rest[:start])
                    r = p.add_run(code.group(1))
                    r.font.name = 'Consolas'
                    r.font.size = Pt(10)
                    rest = rest[end:]
                elif m:
                    start, end = m.span()
                    if start > 0:
                        p.add_run(rest[:start])
                    r = p.add_run(m.group(1))
                    r.bold = True
                    rest = rest[end:]
                else:
                    p.add_run(rest)
                    break
            i += 1
        else:
            i += 1

    if table_lines:
        rows = parse_md_table(table_lines)
        if rows:
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = 'Table Grid'
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    if ci < len(t.rows[ri].cells):
                        t.rows[ri].cells[ci].text = cell

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    backend_dir = os.path.dirname(os.path.abspath(__file__))
    root = os.path.dirname(backend_dir)
    md_path = os.path.join(root, 'HOSTING.md')
    docx_path = os.path.join(root, 'HOSTING.docx')
    if not os.path.exists(md_path):
        print(f"Not found: {md_path}")
        sys.exit(1)
    md_to_docx(md_path, docx_path)
